"""Patch Statistical_Report.docx:

1. Renumber every 'Figure N' where N >= 6 up by 1 (run-level only — never
   overwrites paragraph.text, so manual formatting is preserved).
2. Insert fig_dunn_heatmap.png as new Figure 6 after the KW table in Section 2.4.
"""

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = Path("Statistical_Report.docx")
HEATMAP = Path("figures/fig_dunn_heatmap.png")

CAPTION_TEXT = (
    "Figure 6. Dunn post-hoc Holm-adjusted pairwise significance of profit "
    "distributions across duration bins. Left: hour bins (7×7); right: day "
    "bins (5×5). Blue — Holm-adjusted p < 0.05 (significant); "
    "amber — p ≥ 0.05 (n.s.). Diagonal masked."
)

CAVEAT_TEXT = (
    "At n ≈2.4 M trades statistical power is extreme and almost all "
    "pairs are significant after Holm correction; the informative result is the "
    "two non-significant pairs: 1h–3h vs 8h–9h and 3h–6h vs "
    "6h–8h (Holm p = 0.54), indicating these mid-range hour bins "
    "have indistinguishable profit distributions. All day-bin adjacent pairs are "
    "fully separated."
)


# ---------------------------------------------------------------------------
# 1. Run-level figure renumbering
# ---------------------------------------------------------------------------

def _increment_run_text(text: str) -> str:
    """Increment every 'Figure N' where N >= 6 by 1, in a single regex pass."""
    def _replace(m: re.Match) -> str:
        n = int(m.group(1))
        return f"Figure {n + 1}" if n >= 6 else m.group(0)
    return re.sub(r"Figure (\d+)", _replace, text)


def renumber_figures(doc: Document) -> int:
    """Edit runs in all paragraphs (and table cells); return count of modified runs."""
    count = 0

    def _process_para(para):
        nonlocal count
        for run in para.runs:
            new = _increment_run_text(run.text)
            if new != run.text:
                run.text = new
                count += 1

    for para in doc.paragraphs:
        _process_para(para)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para)
    return count


# ---------------------------------------------------------------------------
# 2. Insert heatmap after the KW table
# ---------------------------------------------------------------------------

def _insert_para_before(doc: Document, ref_para, text: str = "",
                         italic: bool = False, centered: bool = False,
                         space_after_pt: float = 0) -> object:
    """Append a paragraph to the doc then move it before ref_para."""
    new_p = doc.add_paragraph(text)
    if centered:
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text and new_p.runs:
        run = new_p.runs[0]
        run.italic = italic
    if space_after_pt:
        new_p.paragraph_format.space_after = Pt(space_after_pt)
    ref_para._p.addprevious(new_p._p)
    return new_p


def _insert_picture_before(doc: Document, ref_para, image_path: Path,
                             width_inches: float = 6.5) -> object:
    """Append a picture paragraph then move it before ref_para."""
    new_p = doc.add_paragraph()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    ref_para._p.addprevious(new_p._p)
    return new_p


def insert_dunn_figure(doc: Document) -> None:
    """Insert heatmap + caption + caveat before paragraph 56 (the body text
    that follows the KW table empty-spacing paragraph)."""

    # Para[56] in the ORIGINAL document is 'Both tests are highly significant...'
    # After renumbering (run-text only, no structural change) it is still index 56.
    ref = doc.paragraphs[56]

    # Insert in REVERSE order so each addprevious pushes the right element first
    # Final order will be: [img] [caption] [caveat] [ref_para]

    # 3. Caveat sentence (body text, goes last → inserted first)
    caveat_p = _insert_para_before(doc, ref, CAVEAT_TEXT,
                                    italic=False, centered=False, space_after_pt=6)

    # 2. Caption (italic, centered, inserted before caveat → goes before it)
    cap_p = _insert_para_before(doc, caveat_p, CAPTION_TEXT,
                                 italic=True, centered=True, space_after_pt=10)

    # 1. Picture (inserted before caption → ends up first)
    _insert_picture_before(doc, cap_p, HEATMAP, width_inches=6.5)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Loading document…")
    doc = Document(str(DOCX))

    print("Renumbering figures >= 6 at run level...")
    n_runs = renumber_figures(doc)
    print(f"  Modified {n_runs} run(s)")

    # Verify no double-increments: check a sample (ASCII-safe)
    sample = [(i, p.text[:80].encode("ascii", "replace").decode())
              for i, p in enumerate(doc.paragraphs)
              if "Figure" in p.text and any(c.isdigit() for c in p.text)]
    print("  Sample figure references after renumber:")
    for i, t in sample:
        print(f"    [{i}] {t}")

    print("Inserting Dunn heatmap as Figure 6…")
    insert_dunn_figure(doc)

    print("Saving…")
    doc.save(str(DOCX))
    print(f"Saved {DOCX}  ({DOCX.stat().st_size // 1024} KB)")
